from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)
DATE_TEXT = "5 มิถุนายน 2026"
FONT_PATH = r"C:\Windows\Fonts\tahoma.ttf"
BOLD_FONT_PATH = r"C:\Windows\Fonts\tahomabd.ttf"


API_ROUTES = [
    ("Health", "GET", "/api/health", "ตรวจสถานะ backend สำหรับ monitoring", "Public"),
    ("Auth", "POST", "/api/auth/register", "สมัคร user ด้วย email/password", "Public"),
    ("Auth", "POST", "/api/auth/login", "Login admin/email account", "Public"),
    ("Auth", "POST", "/api/auth/password-reset/request", "ขอ reset password", "Public"),
    ("Auth", "POST", "/api/auth/password-reset/confirm", "ยืนยัน reset password", "Public"),
    ("Auth", "POST", "/api/auth/change-password", "เปลี่ยนรหัสผ่านหลัง login", "JWT"),
    ("Auth", "GET", "/api/auth/streamlabs", "สร้าง URL สำหรับ Streamlabs OAuth", "Public"),
    ("Auth", "GET", "/api/auth/streamlabs/connect", "เชื่อมต่อ Streamlabs จาก creator dashboard", "JWT"),
    ("Auth", "GET", "/api/auth/streamlabs/callback", "รับ callback code จาก Streamlabs", "Public"),
    ("Auth", "POST", "/api/auth/streamlabs/exchange", "แลก code/session เป็น token ฝั่ง TipHouse", "Public/JWT flow"),
    ("Public Donation", "GET", "/api/page/:slug", "ดึงข้อมูลหน้าโดเนท public", "Public"),
    ("Public Donation", "GET", "/api/donations/latest/:slug", "ดึงรายการโดเนทล่าสุด", "Public"),
    ("Public Donation", "GET", "/api/donations/rank/:slug", "ดึง Top Donator Rank", "Public"),
    ("Public Donation", "GET", "/api/donations/receipt/:ref", "ดึงใบยืนยันรายการชำระเงินสำเร็จ", "Public"),
    ("Donation", "POST", "/api/donate", "สร้างรายการโดเนท/QR pending", "Public"),
    ("Creator", "GET", "/api/dashboard", "ดึง dashboard creator, history และ stats", "JWT"),
    ("Creator", "PATCH", "/api/settings/page", "บันทึกหน้าตั้งค่าโดเนท", "JWT"),
    ("Settings", "GET", "/api/settings/profile", "ดึงข้อมูลโปรไฟล์ creator", "JWT"),
    ("Settings", "PATCH", "/api/settings/profile", "ส่งคำขอเปลี่ยน profile/email", "JWT"),
    ("Settings", "POST", "/api/settings/onboarding", "บันทึกข้อมูล creator ครั้งแรก", "JWT"),
    ("Settings", "GET", "/api/settings/payout", "ดึงข้อมูลบัญชีโอนจ่าย", "JWT"),
    ("Settings", "PATCH", "/api/settings/payout", "บันทึกข้อมูลบัญชีโอนจ่าย", "JWT"),
    ("Overlay", "GET", "/api/settings/overlay", "ดึงค่าตั้งค่า overlay", "JWT"),
    ("Overlay", "PATCH", "/api/settings/overlay", "บันทึก HTML/CSS/JS, TTS, sound", "JWT"),
    ("Overlay", "POST", "/api/settings/overlay/reset-url", "generate overlay URL ใหม่", "JWT"),
    ("Overlay", "POST", "/api/settings/overlay/test", "ส่ง test alert เข้า OBS overlay", "JWT"),
    ("Overlay Public", "GET", "/api/overlay/tts/th", "TTS endpoint ภาษาไทย", "Public"),
    ("Overlay Public", "GET", "/api/overlay/:key", "ดึง config overlay จาก key", "Public"),
    ("Payment", "POST", "/api/payment/webhook/omise", "รับ webhook จาก Omise", "Gateway"),
    ("Payment", "POST", "/api/payment/webhook/gbprimepay", "รับ webhook จาก GBPrimePay", "Gateway"),
    ("Admin", "GET", "/api/admin/overview", "ข้อมูลสรุป admin", "JWT Admin"),
    ("Admin", "GET", "/api/admin/users", "ค้นหา/แสดงรายการ user", "JWT Admin/Accounting"),
    ("Admin", "PATCH", "/api/admin/users/:id", "แก้ไข user/page/admin data", "JWT Admin"),
    ("Admin", "DELETE", "/api/admin/users/:id", "ลบ user", "JWT Admin"),
    ("Admin", "GET", "/api/admin/transactions", "ค้นหา transaction/donation ทั้งระบบ", "JWT Admin/Accounting"),
    ("Admin", "POST", "/api/admin/transactions/import", "import รายการจาก Excel", "JWT Admin/Accounting"),
    ("Admin", "POST", "/api/admin/transactions/:id/replay-alert", "ส่ง alert ซ้ำไป overlay", "JWT Admin/Accounting"),
    ("Admin", "GET", "/api/admin/transactions/user/:id", "ดู transaction ราย user", "JWT Admin/Accounting"),
    ("Admin", "POST", "/api/admin/users", "สร้าง user/admin/accounting", "JWT Admin"),
    ("Admin", "POST", "/api/admin/users/:id/reset-password", "reset password เป็นค่า default", "JWT Admin"),
    ("Admin", "GET", "/api/admin/approvals", "ดูคำขออนุมัติ", "JWT Admin"),
    ("Admin", "POST", "/api/admin/approvals/:id/approve", "อนุมัติคำขอ", "JWT Admin"),
    ("Admin", "POST", "/api/admin/approvals/:id/reject", "ปฏิเสธคำขอ", "JWT Admin"),
]


def set_run_font(run, size=None, bold=False):
    run.font.name = "Tahoma"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    run._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text, style=None, bold=False):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    return p


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Tahoma"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        style._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")


def add_docx_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            run = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(run, size=9)


def build_docx(path, title, intro, sections):
    doc = Document()
    style_doc(doc)
    add_paragraph(doc, title, "Title", bold=True)
    add_paragraph(doc, intro)
    for section in sections:
        add_paragraph(doc, section["title"], "Heading 1", bold=True)
        for text in section.get("paragraphs", []):
            add_paragraph(doc, text)
        for text in section.get("bullets", []):
            add_paragraph(doc, text, "List Bullet")
        if "table" in section:
            add_docx_table(doc, section["table"][0], section["table"][1])
    doc.save(path)


pdfmetrics.registerFont(TTFont("TipThai", FONT_PATH))
pdfmetrics.registerFont(TTFont("TipThaiBold", BOLD_FONT_PATH))
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="TipTitle", fontName="TipThaiBold", fontSize=18, leading=24, alignment=1))
styles.add(ParagraphStyle(name="TipH1", fontName="TipThaiBold", fontSize=12, leading=16, spaceBefore=8, spaceAfter=4))
styles.add(ParagraphStyle(name="TipBody", fontName="TipThai", fontSize=8.7, leading=12))
styles.add(ParagraphStyle(name="TipSmall", fontName="TipThai", fontSize=6.7, leading=8.5))


def para(text, style="TipBody"):
    return Paragraph(str(text).replace("&", "&amp;"), styles[style])


def add_pdf_table(headers, rows):
    data = [[para(header, "TipSmall") for header in headers]]
    data.extend([[para(value, "TipSmall") for value in row] for row in rows])
    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#94a3b8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def build_pdf(path, title, intro, sections, wide=False):
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4) if wide else A4,
        leftMargin=0.9 * cm,
        rightMargin=0.9 * cm,
        topMargin=0.9 * cm,
        bottomMargin=0.9 * cm,
    )
    story = [para(title, "TipTitle"), Spacer(1, 8), para(intro), Spacer(1, 8)]
    for section in sections:
        story.append(para(section["title"], "TipH1"))
        for text in section.get("paragraphs", []):
            story.append(para(text))
        for text in section.get("bullets", []):
            story.append(para("• " + text))
        if "table" in section:
            story.append(add_pdf_table(section["table"][0], section["table"][1]))
            story.append(Spacer(1, 8))
    doc.build(story)


def write_pair(name, title, intro, sections, wide=False):
    build_docx(OUT / f"{name}.docx", title, intro, sections)
    build_pdf(OUT / f"{name}.pdf", title, intro, sections, wide)


env_rows = [
    ("FRONTEND_URL", "Backend", "URL frontend สำหรับ CORS/Socket เช่น https://tiphouse.com"),
    ("DATABASE_URL", "Backend", "PostgreSQL connection string"),
    ("REDIS_URL", "Backend", "Redis connection string สำหรับ queue/pubsub"),
    ("NEXT_PUBLIC_API_URL", "Frontend", "Backend API URL เช่น https://api.tiphouse.com/api"),
    ("NEXT_PUBLIC_SOCKET_URL", "Frontend", "Backend socket root เช่น https://api.tiphouse.com"),
    ("NEXT_PUBLIC_FRONTEND_URL", "Frontend", "Frontend public URL"),
    ("STREAMLABS_CLIENT_ID/SECRET/REDIRECT_URI", "Backend", "OAuth Streamlabs"),
    ("OMISE_* / GBPRIMEPAY_*", "Backend", "Payment gateway keys และ webhook secret"),
]

write_pair(
    "TipHouse_API_Documentation",
    "TipHouse API Documentation",
    f"สรุป API ตามโค้ด backend ล่าสุด | {DATE_TEXT}",
    [
        {"title": "ภาพรวม", "paragraphs": [f"TipHouse ใช้ NestJS REST API และ Socket.io สำหรับระบบโดเนท, Dashboard, Admin, Payment Webhook และ OBS Overlay จำนวน endpoint ปัจจุบัน {len(API_ROUTES)} เส้น"]},
        {"title": "Environment ที่เกี่ยวข้อง", "table": (["ตัวแปร", "ใช้ที่", "คำอธิบาย"], env_rows)},
        {"title": "รายการ API", "table": (["กลุ่ม", "Method", "Path", "หน้าที่", "สิทธิ์"], API_ROUTES)},
        {"title": "ข้อเสนอเรื่องการลด API", "bullets": [
            "จำนวน API ยังเหมาะสมกับ MVP ที่ต้องแยก Public, Creator, Admin, Payment และ Overlay",
            "สามารถลดได้โดยยุบ dashboard/statistics บางส่วน หรือปิด endpoint legacy หากใช้ Streamlabs-only",
            "ไม่ควรรวม Public, Admin และ Payment Webhook เข้าด้วยกัน เพราะจะกระทบ security และ audit log",
        ]},
    ],
    wide=True,
)

write_pair(
    "TipHouse_Database_Preparation",
    "TipHouse Database Preparation",
    f"เอกสารสำหรับเตรียมฐานข้อมูล PostgreSQL | {DATE_TEXT}",
    [
        {"title": "ฐานข้อมูลที่ระบบใช้", "paragraphs": ["TipHouse ใช้ PostgreSQL และ Prisma ORM. Local development ใช้ postgres:16-alpine และ production ปัจจุบันใช้ PostgreSQL ผ่าน DATABASE_URL"]},
        {"title": "ข้อมูลที่ต้องขอจากผู้ให้บริการ Database", "table": (["รายการ", "ค่าที่ต้องการ/ตัวอย่าง", "หมายเหตุ"], [
            ("Engine", "PostgreSQL 16 หรือใหม่กว่า", "แนะนำให้ใกล้เคียง local"),
            ("Database name", "tiphouse", "ตั้งชื่อได้ตาม provider"),
            ("Username/Password", "strong random password", "ห้ามใช้ password ง่ายใน production"),
            ("Host/Port", "host:5432", "backend ต้องเชื่อมต่อได้"),
            ("SSL mode", "require/prefer", "production ควรใช้ SSL"),
            ("Connection URL", "postgresql://user:pass@host:5432/tiphouse?schema=public", "ใส่ env DATABASE_URL"),
            ("Backup", "Daily backup + PITR ถ้ามี", "สำคัญต่อ transaction"),
            ("Access control", "Private network หรือ allowlist backend IP", "ไม่ควรเปิด public โดยไม่จำเป็น"),
        ])},
        {"title": "Schema หลัก", "bullets": ["User", "DonationPage", "Donation", "OverlaySetting", "PayoutAccount", "ApprovalRequest", "PasswordResetToken", "WebhookLog", "AdminLog"]},
        {"title": "ขั้นตอนเปลี่ยน Database", "bullets": [
            "สร้าง PostgreSQL instance ใหม่และรับ DATABASE_URL",
            "ตั้งค่า DATABASE_URL ใน backend environment",
            "รัน Prisma migration: npm run prisma:migrate -w apps/backend",
            "ถ้ามีข้อมูลเดิม ให้ dump/restore ก่อนเปิด production",
            "ทดสอบ login, settings, donation flow และ admin",
        ]},
    ],
)

write_pair(
    "TipHouse_Server_Preparation",
    "TipHouse Server Preparation",
    f"เอกสารสำหรับเตรียมเซิร์ฟเวอร์ Production | {DATE_TEXT}",
    [
        {"title": "Server ปัจจุบัน", "paragraphs": ["Production ปัจจุบัน deploy เป็น Docker web services สำหรับ tiphouse-backend และ tiphouse-frontend. Local development ใช้ Docker Compose มี PostgreSQL, Redis, Backend, Frontend และ NGINX"]},
        {"title": "Spec แนะนำ", "table": (["รายการ", "MVP", "Production แนะนำ"], [
            ("OS", "Ubuntu 22.04/24.04 LTS", "Ubuntu 24.04 LTS"),
            ("CPU", "2 vCPU", "4 vCPU+"),
            ("RAM", "4 GB", "8 GB+"),
            ("Storage", "40 GB SSD", "80 GB+ SSD + backup"),
            ("Runtime", "Docker + Docker Compose", "Docker + CI/CD + monitoring"),
            ("Network", "80/443 public, 22 restricted", "Cloudflare + firewall + private DB"),
            ("Region", "Singapore/Asia", "เลือกใกล้ไทยที่สุด"),
        ])},
        {"title": "Environment ที่ต้องเตรียม", "table": (["Env", "ตัวอย่าง", "ใช้สำหรับ"], env_rows)},
        {"title": "ขั้นตอนย้าย Server", "bullets": [
            "เตรียม server ใหม่ ติดตั้ง Docker, Compose, firewall และ reverse proxy",
            "ตั้ง DNS ให้ชี้ frontend/backend ไป server ใหม่",
            "ตั้ง env ทั้ง backend และ frontend",
            "ตั้ง SSL และบังคับ HTTPS",
            "Deploy container แล้วรัน migration",
            "อัปเดต Streamlabs Redirect URI และ Payment Webhook URL",
            "ทดสอบ login, donation, webhook, OBS overlay และ admin",
        ]},
        {"title": "Payment/Bank Verification ในอนาคต", "bullets": [
            "ไม่ควรตรวจการโอนด้วยการอ่านอีเมลหรือ LINE เป็นตัวหลัก",
            "ควรใช้ Payment Gateway หรือ Bank API ที่มี dynamic QR, callback/webhook และ payment inquiry",
            "บัญชีกลางควรเป็น merchant/juristic account หรือ gateway settlement",
            "เมื่อ webhook ยืนยัน PAID แล้วจึงส่ง Alert ไป OBS หรือ Streamlabs",
        ]},
    ],
)

print("generated TipHouse docs")
